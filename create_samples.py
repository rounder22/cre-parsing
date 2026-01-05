#!/usr/bin/env python3
"""
Sample data for testing CRE Parser
This script creates test Word and Excel files
"""

from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
import os

def create_sample_word_doc():
    """Create a sample CRE document in Word format"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Commercial Real Estate Investment Proposal', 0)
    
    # Property Details
    doc.add_heading('Property Details', level=1)
    doc.add_paragraph('Property Address: 123 Main Street, New York, NY 10001')
    doc.add_paragraph('Property Type: Mixed-Use Commercial')
    doc.add_paragraph('Square Feet: 50,000')
    doc.add_paragraph('Year Built: 2015')
    doc.add_paragraph('Number of Units: 8')
    doc.add_paragraph('Occupancy Rate: 95%')
    
    # Financial Metrics
    doc.add_heading('Financial Metrics', level=1)
    doc.add_paragraph('Net Operating Income (NOI): $2,500,000')
    doc.add_paragraph('Cap Rate: 6.5%')
    doc.add_paragraph('Purchase Price: $38,462,000')
    doc.add_paragraph('Appraised Value: $40,000,000')
    doc.add_paragraph('Annual Gross Income: $3,200,000')
    doc.add_paragraph('Operating Expenses: $700,000')
    doc.add_paragraph('Debt Service: $1,600,000')
    doc.add_paragraph('DSCR: 1.56')
    doc.add_paragraph('IRR: 12.5%')
    
    # Loan Details
    doc.add_heading('Loan Details', level=1)
    doc.add_paragraph('Loan Amount: $28,000,000')
    doc.add_paragraph('Interest Rate: 4.5%')
    doc.add_paragraph('Loan Term: 10 years')
    doc.add_paragraph('Loan Type: Fixed-Rate Senior Mortgage')
    doc.add_paragraph('Lender: Major Bank Corporation')
    doc.add_paragraph('LTV: 73%')
    
    # Tenant Information
    doc.add_heading('Tenant Information', level=1)
    doc.add_paragraph('Major Tenants: Tech Startup Inc., Coffee Shop Co., Fitness Center LLC')
    doc.add_paragraph('Lease Terms: 5-10 years')
    doc.add_paragraph('Tenant Quality: A+ Credit')
    
    # Market Analysis
    doc.add_heading('Market Analysis', level=1)
    doc.add_paragraph('Market: New York Metropolitan Area')
    doc.add_paragraph('Submarket: Manhattan Central Business District')
    doc.add_paragraph('Comparable Properties: Average cap rate 6.2-6.8%')
    doc.add_paragraph('Market Trends: Strong demand, limited supply')
    
    # Risk Factors
    doc.add_heading('Risk Factors', level=1)
    doc.add_paragraph('Risk Factor 1: Economic downturn affecting tenant cash flows')
    doc.add_paragraph('Risk Factor 2: Key tenant lease expiration in 2027')
    doc.add_paragraph('Risk Factor 3: Rising interest rate environment')
    doc.add_paragraph('Mitigation Strategy: Diversified tenant base and long-term leases')
    
    # Save
    filepath = 'samples/Sample_CRE_Investment.docx'
    os.makedirs('samples', exist_ok=True)
    doc.save(filepath)
    print(f"✓ Created sample Word document: {filepath}")

def create_sample_excel_doc():
    """Create a sample CRE financial model in Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Financial Summary"
    
    # Headers
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    
    headers = ["Metric", "Value", "Unit"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
    
    # Data
    data = [
        ["Property Address", "123 Main Street, NYC", ""],
        ["Square Footage", "50000", "SF"],
        ["Annual Gross Income", "3200000", "$"],
        ["Operating Expenses", "700000", "$"],
        ["NOI", "2500000", "$"],
        ["Purchase Price", "38462000", "$"],
        ["Loan Amount", "28000000", "$"],
        ["Interest Rate", "4.5", "%"],
        ["Loan Term", "10", "years"],
        ["Cap Rate", "6.5", "%"],
        ["DSCR", "1.56", "x"],
        ["Expected IRR", "12.5", "%"],
    ]
    
    for row_idx, row_data in enumerate(data, 2):
        for col_idx, value in enumerate(row_data, 1):
            ws.cell(row=row_idx, column=col_idx, value=value)
    
    # Adjust column widths
    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 10
    
    # Add Tenant Sheet
    tenant_ws = wb.create_sheet("Tenants")
    tenant_headers = ["Tenant Name", "Square Feet", "Annual Rent", "Lease Expiration", "Credit Rating"]
    for col, header in enumerate(tenant_headers, 1):
        cell = tenant_ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
    
    tenant_data = [
        ["Tech Startup Inc.", "20000", "800000", "2029", "A+"],
        ["Coffee Shop Co.", "5000", "150000", "2027", "A"],
        ["Fitness Center LLC", "15000", "500000", "2028", "A-"],
        ["Retail Space", "10000", "250000", "2026", "BBB"],
    ]
    
    for row_idx, row_data in enumerate(tenant_data, 2):
        for col_idx, value in enumerate(row_data, 1):
            tenant_ws.cell(row=row_idx, column=col_idx, value=value)
    
    for col in range(1, 6):
        tenant_ws.column_dimensions[chr(64 + col)].width = 20
    
    # Save
    filepath = 'samples/Sample_Financial_Model.xlsx'
    wb.save(filepath)
    print(f"✓ Created sample Excel document: {filepath}")

if __name__ == '__main__':
    print("Creating sample CRE documents for testing...\n")
    create_sample_word_doc()
    create_sample_excel_doc()
    print("\nSample documents ready for testing!")
