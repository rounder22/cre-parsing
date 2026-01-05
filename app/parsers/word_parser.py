from docx import Document
from typing import Dict, Any, List
from app.parsers.base_parser import BaseParser

class WordParser(BaseParser):
    """Parser for Word (.docx) documents"""
    
    def parse(self) -> Dict[str, Any]:
        """Parse Word document and extract structured data"""
        try:
            doc = Document(self.file_path)
            self.extracted_data = {
                'file_type': 'Word',
                'paragraphs': self._extract_paragraphs(doc),
                'tables': self._extract_tables(doc),
                'text': self.extract_text(),
                'sections': len(doc.sections)
            }
            return self.extracted_data
        except Exception as e:
            raise Exception(f"Error parsing Word document: {str(e)}")
    
    def extract_text(self) -> str:
        """Extract all text from Word document"""
        try:
            doc = Document(self.file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error extracting text from Word document: {str(e)}")
    
    def _extract_paragraphs(self, doc) -> List[Dict[str, Any]]:
        """Extract paragraphs with formatting info"""
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    'text': para.text,
                    'style': para.style.name,
                    'level': para.paragraph_format.outline_level
                })
        return paragraphs
    
    def _extract_tables(self, doc) -> List[List[List[str]]]:
        """Extract tables from Word document"""
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables
